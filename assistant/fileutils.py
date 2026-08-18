"""Text extraction utilities for the AI assistant file reader.

Supported formats:
- PDF       (.pdf)          via pypdf
- Word      (.docx)         via python-docx
- Excel     (.xlsx)         via openpyxl
- Plain text (.txt/.md/.csv) direct read
- Images    (.png/.jpg/.jpeg/.webp) passed to the model as vision input
"""

import mimetypes
import re

from django.conf import settings

KIND_PDF = "pdf"
KIND_DOCX = "docx"
KIND_XLSX = "xlsx"
KIND_TEXT = "text"
KIND_IMAGE = "image"
KIND_OTHER = "other"

IMAGE_EXTENSIONS = ("png", "jpg", "jpeg", "webp")

# Rough limit on the size of text produced from a single sheet.
_XLSX_MAX_ROWS = 4000
_XLSX_MAX_COLS = 80

_CLEAN_RE = re.compile(r"[ \t\u00a0]+")


def _clean_whitespace(text):
    """Collapse runs of spaces and replace stray control chars."""
    if not text:
        return ""
    text = text.replace("\x00", "")
    text = _CLEAN_RE.sub(" ", text)
    return text.strip()


def detect_kind(name):
    """Return the storage kind for an uploaded filename."""
    extension = (name.rsplit(".", 1)[-1].lower()
                 if "." in name else "")

    if extension == "pdf":
        return KIND_PDF
    if extension == "docx":
        return KIND_DOCX
    if extension == "xlsx":
        return KIND_XLSX
    if extension in IMAGE_EXTENSIONS:
        return KIND_IMAGE
    if extension in ("txt", "md", "csv", "log", "json"):
        return KIND_TEXT
    return KIND_OTHER


def is_allowed(name):
    """True if the filename extension is allowed for upload."""
    extension = (name.rsplit(".", 1)[-1].lower()
                 if "." in name else "")
    return extension in settings.ASSISTANT_ALLOWED_EXTENSIONS


def is_image(name):
    return detect_kind(name) == KIND_IMAGE


def is_pdf(file_object):
    """Cheap magic-byte check that the file really is a PDF."""
    head = file_object.read(5)
    file_object.seek(0)
    return head == b"%PDF-"


def human_size(num_bytes):
    """Render a byte count in a human friendly form (Arabic/English agnostic)."""
    value = float(num_bytes)
    for unit in ("B", "KB", "MB", "GB"):
        if value < 1024 or unit == "GB":
            if unit == "B":
                return f"{int(value)} {unit}"
            return f"{value:.1f} {unit}"
        value /= 1024
    return f"{num_bytes} B"


def image_mime(name):
    """Best-effort mime type for an image filename."""
    mime, _ = mimetypes.guess_type(name)
    return mime or "image/png"


def extract_text(file_object, name):
    """Return (text, kind) for an uploaded file.

    The caller is responsible for resetting the file cursor where needed.
    Raises ValueError with a friendly message when the file cannot be read.
    """
    kind = detect_kind(name)

    if kind == KIND_IMAGE:
        return "", KIND_IMAGE

    if kind == KIND_PDF:
        text = _extract_pdf(file_object)
        return _clean_whitespace(text), KIND_PDF

    if kind == KIND_DOCX:
        text = _extract_docx(file_object)
        return _clean_whitespace(text), KIND_DOCX

    if kind == KIND_XLSX:
        text = _extract_xlsx(file_object)
        return _clean_whitespace(text), KIND_XLSX

    if kind == KIND_TEXT:
        text = _extract_text(file_object)
        return _clean_whitespace(text), KIND_TEXT

    raise ValueError(f"unsupported file type: {name}")


def _extract_pdf(file_object):
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ValueError("PDF support is not installed on this server.")

    try:
        reader = PdfReader(file_object)
    except Exception as exc:
        raise ValueError(f"cannot open PDF file: {exc}")

    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(parts)


def _extract_docx(file_object):
    try:
        import docx
    except ImportError:
        raise ValueError("Word document support is not installed.")

    try:
        document = docx.Document(file_object)
    except Exception as exc:
        raise ValueError(f"cannot open Word document: {exc}")

    parts = [p.text for p in document.paragraphs if p.text]

    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_xlsx(file_object):
    try:
        import openpyxl
    except ImportError:
        raise ValueError("Excel support is not installed on this server.")

    try:
        workbook = openpyxl.load_workbook(
            file_object,
            read_only=True,
            data_only=True,
        )
    except Exception as exc:
        raise ValueError(f"cannot open Excel file: {exc}")

    parts = []

    for sheet in workbook.worksheets:
        parts.append(f"### Sheet: {sheet.title} ###")

        for row_index, row in enumerate(sheet.iter_rows(
            values_only=True,
            max_row=_XLSX_MAX_ROWS,
        )):
            if row_index >= _XLSX_MAX_ROWS:
                break

            values = []

            for cell in row[:_XLSX_MAX_COLS]:
                if cell is None:
                    continue
                if isinstance(cell, float) and cell.is_integer():
                    cell = int(cell)
                values.append(str(cell))

            if values:
                parts.append(" | ".join(values))

    workbook.close()
    return "\n".join(parts)


def _extract_text(file_object):
    try:
        data = file_object.read()
    except Exception as exc:
        raise ValueError(f"cannot read text file: {exc}")

    for encoding in ("utf-8-sig", "utf-8", "cp1256", "windows-1252"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue

    return data.decode("utf-8", errors="replace")


def truncate(text, max_chars):
    """Cut text at a word boundary and append an ellipsis."""
    if not text:
        return ""
    if len(text) <= max_chars:
        return text
    cut = text[:max_chars]
    last_space = cut.rfind(" ")
    if last_space > max_chars * 0.6:
        cut = cut[:last_space]
    return cut + "…"
