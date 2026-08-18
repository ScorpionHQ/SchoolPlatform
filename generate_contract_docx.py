"""
إنشاء عقد Word احترافي لمنصة إدارة المدرّسات
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


# Color scheme
GOLD = RGBColor(0xB8, 0x86, 0x0B)
GOLD_LIGHT = RGBColor(0xDA, 0xA5, 0x20)
GOLD_DARK = RGBColor(0x8B, 0x69, 0x14)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY_DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)
GRAY_LIGHT = RGBColor(0x99, 0x99, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CREAM = RGBColor(0xFA, 0xFA, 0xF5)


def set_cell_background(cell, color):
    """Set cell background color. Accepts RGBColor or hex string."""
    if isinstance(color, RGBColor):
        hex_color = str(color)
    else:
        hex_color = color
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.5):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def add_styled_paragraph(doc, text, font_name="Cairo", font_size=12, 
                         bold=False, color=BLACK, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                         rtl=True):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    p.alignment = alignment
    
    # Set RTL
    pPr = p._element.get_or_add_pPr()
    if rtl:
        pPr.set(qn('w:bidi'), '')
    
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = color
    
    return p


def add_section_number(doc, number):
    """Add a section number in gold circle."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Set RTL
    pPr = p._element.get_or_add_pPr()
    pPr.set(qn('w:bidi'), '')
    
    # Add number with formatting
    run = p.add_run(f"  {number}  ")
    run.font.name = "Cairo"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = WHITE
    
    # Background color for the number
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{str(GOLD)}"/>')
    run._element.get_or_add_rPr().append(shading_elm)
    
    return p


def create_contract_table(doc, headers, rows, total_row=None):
    """Create a styled contract table."""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = "Cairo"
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = GOLD_LIGHT
        set_cell_background(cell, BLACK)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Data rows
    for i, row_data in enumerate(rows):
        row = table.rows[i + 1]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_text))
            run.font.name = "Cairo"
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Alternate row colors
            if i % 2 == 0:
                set_cell_background(cell, CREAM)
    
    # Total row
    if total_row:
        total_row_idx = len(rows) + 1
        table.add_row()
        total = table.rows[total_row_idx]
        for j, cell_text in enumerate(total_row):
            cell = total.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_text))
            run.font.name = "Cairo"
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = GOLD_LIGHT
            set_cell_background(cell, BLACK)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    return table


def generate_contract():
    """Generate the professional Word contract."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Cairo'
    font.size = Pt(11)
    font.color.rgb = BLACK
    
    # Set RTL for the document style
    rPr = doc.styles['Normal'].element.get_or_add_rPr()
    rPr.set(qn('w:rtl'), '1')
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ===== HEADER =====
    # Add gold line at top
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 50)
    run.font.color.rgb = GOLD
    run.font.size = Pt(14)
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._element.get_or_add_pPr()
    pPr.set(qn('w:bidi'), '')
    run = p.add_run("عقد اشتراك منصة إدارة المدرّسات")
    run.font.name = "Amiri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = GOLD
    
    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._element.get_or_add_pPr()
    pPr.set(qn('w:bidi'), '')
    run = p.add_run("School Platform Subscription Agreement")
    run.font.name = "Cairo"
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    
    # Gold line at bottom
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 50)
    run.font.color.rgb = GOLD
    run.font.size = Pt(14)
    
    # ===== PARTIES =====
    doc.add_paragraph()  # Spacer
    
    p = add_styled_paragraph(doc, "الأطراف", font_size=14, bold=True, 
                            color=GOLD, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Create parties table
    parties_table = doc.add_table(rows=1, cols=2)
    parties_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # First party
    cell1 = parties_table.rows[0].cells[0]
    cell1.text = ""
    p = cell1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("الطرف الأول (مزوّد الخدمة)")
    run.font.name = "Cairo"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = GOLD
    set_cell_background(cell1, CREAM)
    
    # Add fields for first party
    fields1 = ["الاسم:", "العنوان:", "الهاتف:", "البريد الإلكتروني:"]
    for field in fields1:
        p = cell1.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(field)
        run.font.name = "Cairo"
        run.font.size = Pt(10)
        run.font.color.rgb = GRAY
        p.add_run("\n" + "_" * 30)
    
    # Second party
    cell2 = parties_table.rows[0].cells[1]
    cell2.text = ""
    p = cell2.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("الطرف الثاني (المؤسسة التعليمية)")
    run.font.name = "Cairo"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = GOLD
    set_cell_background(cell2, CREAM)
    
    # Add fields for second party
    fields2 = ["اسم المؤسسة:", "العنوان:", "هاتف المسؤول:", 
               "البريد الإلكتروني:", "كود مدير المؤسسة:"]
    for field in fields2:
        p = cell2.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(field)
        run.font.name = "Cairo"
        run.font.size = Pt(10)
        run.font.color.rgb = GRAY
        p.add_run("\n" + "_" * 30)
    
    # ===== ARTICLE 1: موضوع العقد =====
    doc.add_paragraph()  # Spacer
    
    p = add_styled_paragraph(doc, "المادة الأولى: موضوع العقد", font_size=14, 
                            bold=True, color=BLACK)
    
    p = add_styled_paragraph(doc, "يتعهد الطرف الأول بتقديم منصة إلكترونية لإدارة المدرّسات تشمل:")
    
    items1 = [
        "نظام إدارة المستخدمين (مدير / معلم / طالب / ولي أمر)",
        "نظام تسجيل الحضور والغياب",
        "نظام إدخال الدرجات والتقارير الأكاديمية",
        "نظام الفواتير والمدفوعات",
        "نظام الاستيراد الجماعي للطلاب",
        "المساعد الذكي للاستفسارات",
        "تطبيق ويب قابل للتثبيت (PWA)",
        "دعم كامل للغة العربية واجهة من اليمين لليسار"
    ]
    
    for i, item in enumerate(items1, 1):
        p = add_styled_paragraph(doc, f"{i}. {item}")
    
    # ===== ARTICLE 2: مدة العقد =====
    doc.add_paragraph()  # Spacer
    
    p = add_styled_paragraph(doc, "المادة الثانية: مدة العقد", font_size=14, 
                            bold=True, color=BLACK)
    
    p = add_styled_paragraph(doc, 
        "مدة العقد: سنة واحدة (12 شهراً) تبدأ من تاريخ التفعيل")
    p = add_styled_paragraph(doc, 
        "يتجدد العقد تلقائياً لمدة مماثلة ما لم يُبلغ أحد الطرفين الآخر برغبته في الإلغاء قبل 30 يوماً من انتهاء المدة")
    
    # ===== ARTICLE 3: الباقة والسعر =====
    doc.add_paragraph()  # Spacer
    
    p = add_styled_paragraph(doc, "المادة الثالثة: الباقة والسعر", font_size=14, 
                            bold=True, color=BLACK)
    
    p = add_styled_paragraph(doc, "الباقة الشاملة", font_size=12, bold=True, color=GOLD)
    
    # Package table
    package_headers = ["البند", "الحد"]
    package_rows = [
        ["عدد الطلاب", "غير محدود"],
        ["عدد المعلمين", "غير محدود"],
        ["عدد الصفوف", "غير محدود"],
        ["التخزين", "50 GB"],
        ["الدعم الفني", "هاتف + واتساب + بريد"],
        ["المساعد الذكي", "مشمول"],
        ["التقارير المتقدمة", "مشمولة"],
        ["تحليلات الذكاء الاصطناعي", "مشمولة"],
        ["نسخ احتياطي يومي", "مشمول"]
    ]
    
    create_contract_table(doc, package_headers, package_rows)
    
    p = add_styled_paragraph(doc, "السعر السنوي: 7,000,000 دينار عراقي", 
                            font_size=14, bold=True, color=GOLD_DARK,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    # ===== ARTICLE 4: الأموال التشغيلية =====
    doc.add_paragraph()  # Spacer
    
    p = add_styled_paragraph(doc, "المادة الرابعة: الأموال التشغيلية", font_size=14, 
                            bold=True, color=BLACK)
    
    p = add_styled_paragraph(doc, 
        "تُضاف إلى سعر الاشتراك سنوياً مبلغ 1,500,000 دينار عراقي كأجر تشغيلي يغطي التكاليف التالية:")
    
    # Operating costs table
    costs_headers = ["البند", "التكلفة السنوية", "التفاصيل"]
    costs_rows = [
        ["الاستضافة والخادم", "500,000 دينار", "خادم افتراضي (4GB ذاكرة، 2 معالج، 100GB تخزين) مع نسخ احتياطي تلقائي"],
        ["النطاق والشهادة الأمنية", "100,000 دينار", "تسجيل اسم النطاق + شهادة التشفير لمدة سنة كاملة"],
        ["الصيانة والتحديثات", "400,000 دينار", "تطوير مستمر للمنصة + إصلاح الأخطاء البرمجية + إضافة ميزات جديدة"],
        ["الدعم الفني", "300,000 دينار", "فريق دعم فني بدوام جزئي (8 ساعات / أسبوع)"],
        ["الأمان والحماية", "100,000 دينار", "جدار حماية + مراقبة أمنية مستمرة + حماية من الاختراق"],
        ["النسخ الاحتياطي والاستعادة", "100,000 دينار", "نسخ احتياطي يومي للبيانات + استعادة في حالات الكوارث"],
        ["الكهرباء والاتصالات", "100,000 دينار", "استهلاك الكهرباء + اتصال الإنترنت"]
    ]
    
    create_contract_table(doc, costs_headers, costs_rows, total_row=["الإجمالي", "1,500,000 دينار عراقي", ""])
    
    # Notes
    p = add_styled_paragraph(doc, "ملاحظات:", font_size=11, bold=True, color=GOLD_DARK)
    notes = [
        "تُدفع هذه الرسوم سنوياً مع اشتراك الباقة",
        "تُذكر في فاتورة منفصلة أو كجزء من الفاتورة الرئيسية",
        "لا تُسترجع في حال إلغاء العقد",
        "قد تتغير حسب معدل التضخم وتغيرات السوق"
    ]
    for note in notes:
        p = add_styled_paragraph(doc, f"• {note}", font_size=10)
    
    # ===== ARTICLE 5: التكاليف الإضافية =====
    doc.add_paragraph()  # Spacer
    
    p = add_styled_paragraph(doc, "المادة الخامسة: التكاليف الإضافية", font_size=14, 
                            bold=True, color=BLACK)
    
    # Additional costs table
    additional_headers = ["البند", "السعر"]
    additional_rows = [
        ["طالب إضافي", "10,000 دينار / شهر"],
        ["معلم إضافي", "25,000 دينار / شهر"],
        ["تدريب إضافي (يوم واحد)", "100,000 دينار"],
        ["تخصيص واجهة (شعار + ألوان)", "500,000 دينار (مرة واحدة)"]
    ]
    
    create_contract_table(doc, additional_headers, additional_rows)
    
    # ===== ARTICLE 6: شروط الدفع =====
    doc.add_paragraph()  # Spacer
    
    p = add_styled_paragraph(doc, "المادة السادسة: شروط الدفع", font_size=14, 
                            bold=True, color=BLACK)
    
    p = add_styled_paragraph(doc, "ملخص المبالغ", font_size=12, bold=True, color=GOLD)
    
    # Summary table
    summary_headers = ["البند", "المبلغ"]
    summary_rows = [
        ["اشتراك الباقة الشاملة (سنوياً)", "7,000,000 دينار عراقي"],
        ["الأموال التشغيلية (سنوياً)", "1,500,000 دينار عراقي"]
    ]
    
    create_contract_table(doc, summary_headers, summary_rows, total_row=["الإجمالي السنوي", "8,500,000 دينار عراقي"])
    
    p = add_styled_paragraph(doc, "طرق الدفع", font_size=12, bold=True, color=GOLD)
    
    # Payment methods table
    payment_headers = ["مرحلة الدفع", "النسبة", "المبلغ"]
    payment_rows = [
        ["الدفعة المقدمة (قبل التوقيع)", "30%", "2,550,000 دينار عراقي"],
        ["الدفعة المتبقية (بعد التوقيع)", "70%", "5,950,000 دينار عراقي"]
    ]
    
    create_contract_table(doc, payment_headers, payment_rows)
    
    p = add_styled_paragraph(doc, "الدفع السنوي فقط", font_size=11, bold=True)
    p = add_styled_paragraph(doc, "تُقبل الطرق التالية: نقدي / تحويل بنكي / شيك")
    p = add_styled_paragraph(doc, "تُصدر الفواتير رسمياً بتاريخ الاستحقاق")
    
    # ===== ARTICLE 7: الملكية والبيانات =====
    doc.add_paragraph()  # Spacer
    
    p = add_styled_paragraph(doc, "المادة السابعة: الملكية والبيانات", font_size=14, 
                            bold=True, color=BLACK)
    
    items7 = [
        "جميع بيانات المؤسسة ملكيتها الحصرية",
        "يلتزم الطرف الأول بحماية البيانات وخصوصيتها",
        "تصدير جميع البيانات عند انتهاء العقد مجاناً",
        "حذف البيانات من خوادم الطرف الأول خلال 30 يوماً من انتهاء العقد",
        "لا يحق للطرف الأول استخدام البيانات لأي غرض آخر"
    ]
    
    for i, item in enumerate(items7, 1):
        p = add_styled_paragraph(doc, f"{i}. {item}")
    
    # ===== ARTICLE 8: الدعم الفني والصيانة =====
    doc.add_paragraph()  # Spacer
    
    p = add_styled_paragraph(doc, "المادة الثامنة: الدعم الفني والصيانة", font_size=14, 
                            bold=True, color=BLACK)
    
    items8 = [
        "تحديثات مستمرة مجاناً طوال مدة العقد",
        "إصلاح الأخطاء البرمجية مجاناً",
        "الدعم الفني: 8 ساعات / أسبوع، وقت الاستجابة: 48 ساعة كحد أقصى",
        "التدريب الأولي: يوم واحد مجاني (4 ساعات)",
        "التدريب الإضافي: 100,000 دينار / يوم"
    ]
    
    for i, item in enumerate(items8, 1):
        p = add_styled_paragraph(doc, f"{i}. {item}")
    
    # ===== ARTICLE 9: الإلغاء والاسترداد =====
    doc.add_paragraph()  # Spacer
    
    p = add_styled_paragraph(doc, "المادة التاسعة: الإلغاء والاسترداد", font_size=14, 
                            bold=True, color=BLACK)
    
    items9 = [
        "يحق للطرف الثاني الإلغاء مع إشعار كتابي قبل 30 يوماً",
        "في حال الإلغاء قبل 6 أشهر: استرداد 50% من المبلغ المدفوع",
        "في حال الإلغاء بعد 6 أشهر: لا يستحق استرداد",
        "يحق للطرف الأول تعليق الخدمة في حال عدم الدفع لمدة 15 يوماً بعد موعد الاستحقاق"
    ]
    
    for i, item in enumerate(items9, 1):
        p = add_styled_paragraph(doc, f"{i}. {item}")
    
    # ===== ARTICLE 10: المسؤولية =====
    doc.add_paragraph()  # Spacer
    
    p = add_styled_paragraph(doc, "المادة العاشرة: المسؤولية", font_size=14, 
                            bold=True, color=BLACK)
    
    items10 = [
        "يلتزم الطرف الأول بتوفير الخدمة بنسبة 99% وقت التشغيل شهرياً",
        "في حال تجاوز التوقف عن العمل ساعتين متتاليتين، يُخصم من الفاتورة التالية",
        "لا يتحمل الطرف الأول المسؤولية عن الأخطاء الناتجة عن سوء الاستخدام",
        "تُحل النزاعات ودياً أولاً، ثم عبر الجهات القضائية المختصة"
    ]
    
    for i, item in enumerate(items10, 1):
        p = add_styled_paragraph(doc, f"{i}. {item}")
    
    # ===== ARTICLE 11: السرية =====
    doc.add_paragraph()  # Spacer
    
    p = add_styled_paragraph(doc, "المادة الحادية عشرة: السرية", font_size=14, 
                            bold=True, color=BLACK)
    
    items11 = [
        "يلتزم كل طرف بالحفاظ على سرية المعلومات الخاصة بالطرف الآخر",
        "لا يُفصح عن أي معلومات دون موافقة كتابية مسبقة",
        "تستمر التزامات السرية لمدة سنتين بعد انتهاء العقد"
    ]
    
    for i, item in enumerate(items11, 1):
        p = add_styled_paragraph(doc, f"{i}. {item}")
    
    # ===== ARTICLE 12: القوة القاهرة =====
    doc.add_paragraph()  # Spacer
    
    p = add_styled_paragraph(doc, "المادة الثانية عشرة: القوة القاهرة", font_size=14, 
                            bold=True, color=BLACK)
    
    p = add_styled_paragraph(doc, 
        "لا يُعتبر أي طرف مخلاً بالعقد إذا كان عدم التنفيذ ناتجاً عن قوة قاهرة كالأوبئة، الكوارث الطبيعية، الحروب، أو التعطل الشبكي كبير الحجم.")
    
    # ===== ARTICLE 13: التعديلات =====
    doc.add_paragraph()  # Spacer
    
    p = add_styled_paragraph(doc, "المادة الثالثة عشرة: التعديلات", font_size=14, 
                            bold=True, color=BLACK)
    
    items13 = [
        "يُستثنى من هذا العقد أي تعديلات يتم الاتفاق عليها كتابياً بين الطرفين",
        "أي تعديلات يجب أن يوقعها الطرفان"
    ]
    
    for i, item in enumerate(items13, 1):
        p = add_styled_paragraph(doc, f"{i}. {item}")
    
    # ===== ARTICLE 14: القانون الحاكم =====
    doc.add_paragraph()  # Spacer
    
    p = add_styled_paragraph(doc, "المادة الرابعة عشرة: القانون الحاكم والاختصاص القضائي", 
                            font_size=14, bold=True, color=BLACK)
    
    items14 = [
        "يخضع هذا العقد لقوانين الجمهورية العراقية",
        "أي نزاع يُحل عبر المحاكم المختصة في محافظة نينوى"
    ]
    
    for i, item in enumerate(items14, 1):
        p = add_styled_paragraph(doc, f"{i}. {item}")
    
    # ===== SIGNATURES =====
    doc.add_paragraph()  # Spacer
    
    p = add_styled_paragraph(doc, "التوقيعات", font_size=16, bold=True, 
                            color=BLACK, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Signatures table
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # First party signature
    cell1 = sig_table.rows[0].cells[0]
    cell1.text = ""
    p = cell1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("الطرف الأول (مزوّد الخدمة)")
    run.font.name = "Cairo"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = GOLD
    set_cell_background(cell1, CREAM)
    
    p = cell1.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n" + "_" * 30)
    run.font.name = "Cairo"
    run.font.size = Pt(10)
    
    p = cell1.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("الاسم والتوقيع")
    run.font.name = "Cairo"
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    
    p = cell1.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n" + "_" * 30)
    run.font.name = "Cairo"
    run.font.size = Pt(10)
    
    p = cell1.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("التاريخ: ___/___/______")
    run.font.name = "Cairo"
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    
    p = cell1.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n" + "_" * 30)
    run.font.name = "Cairo"
    run.font.size = Pt(10)
    
    p = cell1.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("الختم")
    run.font.name = "Cairo"
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    
    # Second party signature
    cell2 = sig_table.rows[0].cells[1]
    cell2.text = ""
    p = cell2.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("الطرف الثاني (المؤسسة التعليمية)")
    run.font.name = "Cairo"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = GOLD
    set_cell_background(cell2, CREAM)
    
    p = cell2.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n" + "_" * 30)
    run.font.name = "Cairo"
    run.font.size = Pt(10)
    
    p = cell2.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("الاسم والتوقيع")
    run.font.name = "Cairo"
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    
    p = cell2.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n" + "_" * 30)
    run.font.name = "Cairo"
    run.font.size = Pt(10)
    
    p = cell2.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("التاريخ: ___/___/______")
    run.font.name = "Cairo"
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    
    p = cell2.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n" + "_" * 30)
    run.font.name = "Cairo"
    run.font.size = Pt(10)
    
    p = cell2.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("الختم")
    run.font.name = "Cairo"
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    
    # ===== FOOTER =====
    doc.add_paragraph()  # Spacer
    
    p = add_styled_paragraph(doc, "تم إعداد هذا العقد بتاريخ: ___/___/______", 
                            font_size=10, color=GRAY,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    p = add_styled_paragraph(doc, "✦ منصة إدارة المدرّسات ✦", 
                            font_size=12, bold=True, color=GOLD,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Save document
    output_path = r"E:\Desktop\SchoolPlatform\عقد_المنصة_المدرسية.docx"
    doc.save(output_path)
    print("Done: Contract Word document created successfully")
    
    return output_path


if __name__ == "__main__":
    generate_contract()
